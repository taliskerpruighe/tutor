#!/usr/bin/env python3
"""Build the loose docx components for the Stavros Daphne naturalization packet."""
import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = "/home/talisker/garage/tools/tutor/.claude/worktrees/challenges-spike-002/lab/synthetic/dogfood/stavros_daphne/build/docx"
os.makedirs(OUT, exist_ok=True)

PAGE_W = Emu(7772400)
PAGE_H = Emu(10058400)
MARGIN = Emu(914400)

def new_doc():
    d = docx.Document()
    d.styles['Normal'].font.name = 'Times New Roman'
    d.styles['Normal'].font.size = Pt(12)
    pf = d.styles['Normal'].paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    sec = d.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.top_margin = MARGIN
    sec.bottom_margin = MARGIN
    sec.left_margin = MARGIN
    sec.right_margin = MARGIN
    return d

def add_p(d, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=None, tabs=False):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.bold = bold
        if size:
            r.font.size = Pt(size)
    return p

# ---------------------------------------------------------------------------
# 00. Applicant Cover Page
# ---------------------------------------------------------------------------
d = new_doc()
add_p(d, "APPLICATION FOR NATURALIZATION", align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(d)
add_p(d)
add_p(d, "APPLICANT:")
add_p(d)
add_p(d, "Mr. Konstantinos Stavros")
add_p(d, "DOB: 04/12/1957")
add_p(d, "COB/CON: Greece")
add_p(d)
add_p(d, "Classification Basis: INA 316(a)")
d.save(f"{OUT}/00. Applicant Cover Page.docx")

# ---------------------------------------------------------------------------
# Tab A-0 cover
# ---------------------------------------------------------------------------
d = new_doc()
p = add_p(d, "TAB A", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
p2 = add_p(d, "SUMMARY", align=WD_ALIGN_PARAGRAPH.CENTER, size=24)
d.save(f"{OUT}/A-0. Tab Cover Page.docx")

# ---------------------------------------------------------------------------
# Tab B-0 cover
# ---------------------------------------------------------------------------
d = new_doc()
add_p(d, "TAB B", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
add_p(d, "BIOGRAPHICAL INFORMATION", align=WD_ALIGN_PARAGRAPH.CENTER, size=24)
d.save(f"{OUT}/B-0. Tab Cover Page.docx")

# ---------------------------------------------------------------------------
# A-1 Table of Contents
# ---------------------------------------------------------------------------
d = new_doc()
add_p(d, "TABLE OF CONTENTS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_p(d, "Tab A (Summary)", bold=True)
add_p(d, "1.\tTable of contents")
add_p(d, "2.\tCover letter")
add_p(d, "Tab B (Biographical Information)", bold=True)
add_p(d, "3.\tForm N-400, Application for Naturalization")
add_p(d, "4.\tBio page of latest passport of the applicant")
add_p(d, "5.\tForm I-551, Permanent Resident Card")
add_p(d, "6.\tLatest tax return")
add_p(d, "7.\tTravel addendum")
add_p(d, "8.\tWritten explanation")
d.save(f"{OUT}/A-1. Table of Contents.docx")

# ---------------------------------------------------------------------------
# A-2 Cover Letter
# ---------------------------------------------------------------------------
d = new_doc()
add_p(d, "March 18, 2026")
add_p(d)
add_p(d, "VIA U.S. POSTAL SERVICE (USPS)", bold=True)
add_p(d, "U.S. Department of Homeland Security")
add_p(d, "United States Citizenship and Immigration Services")
add_p(d, "USCIS Elgin Lockbox")
add_p(d, "Attn: N-400")
add_p(d, "P.O. Box 4060")
add_p(d, "Carol Stream, IL 60197-4060")
add_p(d)
add_p(d, "\tRe:\tForm N-400, Application for Naturalization")
add_p(d)
add_p(d, "\t\tApplicant:\tMr. Konstantinos Stavros")
add_p(d, "\t\tDOB:\t\tApril 12, 1957")
add_p(d, "\t\tCOB/CON:\tGreece")
add_p(d)
add_p(d, "To Whom It May Concern:")
add_p(d)
add_p(d, "Enclosed please find one (1) Form N-400, Application for Naturalization, along with the "
         "accompanying filing fee of $760.00 and supporting documentation, for Mr. Konstantinos "
         "Stavros, a Greek national. Mr. Stavros is eligible for naturalization as it has been more "
         "than five (5) years since he became a permanent resident on January 23, 2020. See INA § "
         "316(a); 8 C.F.R. § 316.2.")
add_p(d)
add_p(d, "All supporting documents in this packet are photocopies of originals. Mr. Stavros "
         "understands that he may have to present originals as part of the adjudication process.")
add_p(d)
add_p(d, "We look forward to your speedy and favorable adjudication of this application.")
add_p(d)
add_p(d, "Sincerely,")
add_p(d)
add_p(d)
add_p(d, "Petition Preparer")
d.save(f"{OUT}/A-2. Cover Letter.docx")

# ---------------------------------------------------------------------------
# B-7 Travel Addendum
# ---------------------------------------------------------------------------
trips = [
    ("11/02/2025", "11/20/2025", "Greece"),
    ("04/10/2025", "04/28/2025", "Greece"),
    ("09/05/2024", "09/22/2024", "Greece"),
    ("01/15/2024", "01/30/2024", "Greece"),
    ("06/01/2023", "06/20/2023", "Greece"),
    ("10/10/2022", "10/25/2022", "Greece"),
    ("03/01/2022", "03/15/2022", "Greece"),
    ("08/01/2021", "08/18/2021", "Greece"),
]
d = new_doc()
add_p(d, "TRAVEL ADDENDUM", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_p(d, "The following is a full list of the Applicant's trips to countries other than the United "
         "States within the last 5 years, excluding day trips. It combines the trips listed in Page "
         "6, Part 8, Question 1, with the trips listed in the addendum thereto.")
for i, (dep, ret, country) in enumerate(trips, start=1):
    add_p(d, f"{i}. {dep}-{ret} – {country}")
d.save(f"{OUT}/B-7. Travel Addendum.docx")

# ---------------------------------------------------------------------------
# B-8 Written Explanation
# ---------------------------------------------------------------------------
d = new_doc()
add_p(d, "WRITTEN EXPLANATION", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_p(d, "The following explanation is submitted on behalf of Mr. Konstantinos Stavros in support of "
         "his Form N-400, Application for Naturalization, and responds to the items indicated in "
         "Part 9. Additional Information About You.")
add_p(d, "Part 9., Item Number 8.a.", bold=True)
add_p(d, "Served in the Hellenic Army (Greece), Infantry Branch, with the rank of Private, from "
         "March 1976 to February 1978, pursuant to compulsory national military service required of "
         "Greek men of his generation. He was honorably discharged at the end of his term of service. "
         "This was routine national conscription; he held no rank of command, saw no combat, and had "
         "no involvement with any unit engaged in persecution, human rights abuses, or hostilities of "
         "any kind.")
add_p(d, "Part 9., Item Number 12.", bold=True)
add_p(d, "In connection with the compulsory military service described above at Item Number 8.a., he "
         "received ordinary weapons handling and basic infantry training, standard for every conscript "
         "in his cohort and not exceeding the routine training given to any other young man of his "
         "generation completing Greek national service.")
add_p(d, "Part 9., Item Number 20.", bold=True)
add_p(d, "Placed in removal proceedings on 2015-09-14, following the denial on 2015-02-11 of a Form "
         "I-485, Application to Register Permanent Residence, filed concurrently with a Form I-130, "
         "Petition for Alien Relative, submitted on his behalf by his U.S. citizen daughter on "
         "2013-06-18 (matter referenced by A-Number A-207115634). The Request for Evidence and the "
         "denial notice were both sent to an address he had already moved away from and were never "
         "received; he did not respond because he was never made aware of either notice. He was never "
         "arrested, detained, or charged with any offense in connection with this matter. On "
         "2018-11-06, the removal proceedings were terminated with no order of removal entered against "
         "him; he was never removed or deported from the United States. The Form I-130 petition was "
         "re-filed and later approved, and he was granted lawful permanent residence on 2020-01-23.")
d.save(f"{OUT}/B-8. Written Explanation.docx")

print("done")

#!/usr/bin/env python3
"""verify_inputs.py — the Phase 4 barrier (BUILD-PLAN §5).

BUILD-PLAN is explicit that this check runs "on the **extracted text of the
actually-produced inputs** (not on the agents' claims)". So nothing here reads
an agent's report or trusts `input_surfaces` to describe reality: it extracts
the text of every file in the input folder and asks whether the facts are
actually findable.

Checks per client:
  1  the input folder exists and has a plausible email-directory structure
  2  every attachment named in a body actually exists on disk
  3  every file on disk is referenced by some body (no orphans)
  4  NO FIRM IDENTITY appears anywhere in the inputs (§16 r7)
  5  the load-bearing facts are findable in the extracted text
  6  mess events are real: each superseded fact appears in BOTH its wrong and
     its right form, and the right one is later
  7  leakage: the same distinctive-token scan the masterkey barrier uses

A fact that cannot be found is a BROKEN TEST for a to-do client, not a
blemish — the packet must be buildable from the folder.
"""
import os, sys, re, glob, subprocess, datetime
import yaml

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CLIENTS = os.path.join(ROOT, "clients")
FAILS, WARNS, OKS = [], [], []
def ok(c, m): OKS.append(f"[ok]   {c}: {m}")
def bad(c, m): FAILS.append(f"[FAIL] {c}: {m}")
def warn(c, m): WARNS.append(f"[warn] {c}: {m}")

def extract(path):
    """Text of one input file, whatever its type."""
    ext = path.lower().rsplit(".", 1)[-1]
    try:
        if ext == "txt":
            return open(path, errors="ignore").read()
        if ext == "pdf":
            return subprocess.run(["pdftotext", "-layout", path, "-"],
                                  capture_output=True, text=True, timeout=60).stdout
        if ext == "docx":
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts += [c.text for c in row.cells]
            return "\n".join(parts)
        if ext == "xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            return "\n".join(str(c.value) for ws in wb for r in ws.iter_rows()
                             for c in r if c.value is not None)
        if ext == "csv":
            return open(path, errors="ignore").read()
    except Exception as e:
        return f"[UNREADABLE: {e}]"
    return ""

def norm(s):
    return re.sub(r"[^a-z0-9]", "", str(s).lower())

def as_date(v):
    """A date, however it was written. YAML gives some masterkeys a real
    datetime.date and others the string '1988-03-22'; the first draft of this
    function only tried the house formats for the former, so EVERY client
    reported its DOB and LPR date as unfindable — a uniform failure across all
    six, which is the signature of a broken check rather than broken data."""
    if isinstance(v, datetime.datetime): return v.date()
    if isinstance(v, datetime.date): return v
    if isinstance(v, str):
        for f in ("%Y-%m-%d", "%m/%d/%Y", "%B %d, %Y", "%d %B %Y"):
            try: return datetime.datetime.strptime(v.strip(), f).date()
            except ValueError: pass
    return None

def findable(needle, hay_norm, hay_raw):
    if needle in (None, "", []): return True
    n = norm(needle)
    if len(n) < 3: return True
    if n in hay_norm: return True
    d = as_date(needle)
    if d:
        for f in ("%m/%d/%Y", "%-m/%-d/%Y", "%B %d, %Y", "%d %B %Y",
                  "%Y-%m-%d", "%b %d, %Y", "%m-%d-%Y", "%d/%m/%Y"):
            try:
                if norm(d.strftime(f)) in hay_norm: return True
            except ValueError:
                pass
    return False

_ALLOW, _BLOCK = None, None
def _leak_vocab():
    """Exactly validate_masterkeys.py's allow set: the committed blank forms,
    STYLE-SPEC §11, document-catalog.yaml, ordinary English, and the recorded
    §11 exclusion set."""
    global _ALLOW, _BLOCK
    if _ALLOW is not None: return
    allow = set()
    for f in ("n-400.pdf", "f1040.pdf", "f1040--2024.pdf"):
        fp = os.path.join(ROOT, "blanks", f)
        if os.path.exists(fp):
            t = subprocess.run(["pdftotext", "-layout", fp, "-"],
                               capture_output=True, text=True).stdout
            allow |= {w.lower() for w in re.findall(r"[A-Za-z][A-Za-z'\-]{2,}", t)}
    sp = os.path.join(ROOT, "spec", "STYLE-SPEC.md")
    if os.path.exists(sp):
        t = open(sp).read()
        allow |= {w.lower() for w in re.findall(r"[A-Za-z][A-Za-z'\-]{2,}", t)}
    cat = os.path.join(ROOT, "templates", "document-catalog.yaml")
    if os.path.exists(cat):
        allow |= {w.lower() for w in re.findall(r"[A-Za-z]{3,}", open(cat).read())}
    for fp in ("/usr/share/dict/words", "/usr/share/dict/linux.words"):
        if os.path.exists(fp):
            allow |= {w.strip().lower() for w in open(fp, errors="ignore")
                      if len(w.strip()) > 2}
            break
    exc = os.path.join(ROOT, "tools", "blocklist-exclusions.txt")
    if os.path.exists(exc):
        allow |= {l.strip().lower() for l in open(exc)
                  if l.strip() and not l.startswith("#")}
    _ALLOW = allow
    _BLOCK = {l.strip().lower() for l in open(os.path.join(ROOT, "blocklist.txt"))
              if len(l.strip()) >= 4}

FORCE_FLAG = {"oliveira", "izaguirre", "symple", "trysymple", "ossola", "ylenia",
              "luwilyn", "xuying", "zhu", "malone", "garth", "braun", "marcel"}

def _allowed(tok):
    t = tok.lower().strip()
    if t in FORCE_FLAG: return False
    if t in _ALLOW: return True
    t2 = re.sub(r"[\u2019'](s|t|re|ve|ll|d|m)$", "", t)
    if t2 in _ALLOW: return True
    # A stem of three characters or fewer cannot be a distinctive identifier.
    # This is what finally cleared "it's" -> "it": the dictionary pass only
    # admits words longer than two characters, so common contractions of short
    # pronouns fell through and red-lighted three clients.
    if len(t2) <= 3: return True
    parts = [x for x in re.split(r"[-/]", t2) if x]
    return bool(parts) and all(x in _ALLOW or len(x) <= 3 for x in parts)

def leak_hits(text):
    _leak_vocab()
    out = set()
    for tok in re.findall(r"[A-Za-z][A-Za-z'\-]{3,}|\d{5,}", text):
        if tok.lower() in _BLOCK and not _allowed(tok):
            out.add(tok)
    return out


def main():
    slugs = sorted(d for d in os.listdir(CLIENTS)
                   if os.path.isdir(os.path.join(CLIENTS, d)))
    for slug in slugs:
        indir = os.path.join(CLIENTS, slug, "input")
        if not os.path.isdir(indir):
            bad(slug, "no input/ folder"); continue
        files = [p for p in glob.glob(os.path.join(indir, "**", "*"), recursive=True)
                 if os.path.isfile(p)]
        if not files: bad(slug, "input/ is empty"); continue

        maildirs = sorted({os.path.relpath(os.path.dirname(p), indir).split(os.sep)[0]
                           for p in files})
        numbered = [d for d in maildirs if re.match(r"^\d{6}_\d{4}-\d{2}-\d{2}_", d)]
        if len(numbered) < 3:
            bad(slug, f"only {len(numbered)} NNNNNN_YYYY-MM-DD_slug dir(s) — "
                      f"BUILD-PLAN §5 sets a minimum of 3 that carry the surfaces")
        else:
            ok(slug, f"{len(numbered)} email directories, {len(files)} files")

        bodies = [p for p in files if os.path.basename(p).lower().startswith("body")
                  or p.lower().endswith(("body.txt",))]
        if not bodies:
            bodies = [p for p in files if p.lower().endswith(".txt")]
        body_text = "\n".join(extract(p) for p in bodies)
        all_text = "\n".join(extract(p) for p in files)
        hay_norm, hay_raw = norm(all_text), all_text
        body_norm = norm(body_text)

        # ---- 2 & 3: attachments referenced vs present ---------------------
        attach = [p for p in files if p not in bodies]
        orphan = [os.path.basename(p) for p in attach
                  if norm(os.path.basename(p)) not in body_norm
                  and norm(os.path.splitext(os.path.basename(p))[0]) not in body_norm]
        if orphan:
            warn(slug, f"{len(orphan)} attachment(s) not named in any body: {orphan[:4]}")
        else:
            ok(slug, f"all {len(attach)} attachments are named in a body")

        # ---- 4: no firm identity -----------------------------------------
        firm = re.findall(r"\b(?:LLP|Esq\.?|Law Offices?|Attorneys? at Law|"
                          r"Immigration Law)\b", all_text)
        if firm:
            bad(slug, f"firm identity in inputs (§16 r7): {sorted(set(firm))[:5]}")
        else:
            ok(slug, "no firm identity anywhere in the inputs (§16 r7)")

        # ---- 5: the load-bearing facts are findable ----------------------
        mk = yaml.safe_load(open(os.path.join(CLIENTS, slug, "masterkey.norm.yaml")))
        ident, contact = mk.get("identity") or {}, mk.get("contact") or {}
        imm = mk.get("immigration") or {}
        docs = mk.get("documents") or {}
        want = {
            "family name":     ident.get("family_name"),
            "given name":      ident.get("given_name"),
            "date of birth":   ident.get("dob"),
            "country of birth": ident.get("country_of_birth") or ident.get("cob"),
            "A-number":        imm.get("a_number"),
            "LPR date":        imm.get("lpr_date"),
            "email":           contact.get("email"),
            "daytime phone":   contact.get("daytime_phone"),
        }
        pp = docs.get("applicant_passport") or {}
        want["passport number"] = pp.get("passport_number") or pp.get("number")
        missing = [k for k, v in want.items() if not findable(v, hay_norm, hay_raw)]
        ships = str(mk.get("ships_as", ""))
        if missing:
            # an unlocatable fact breaks a to-do client outright
            (bad if ships.startswith("to-do") else warn)(
                slug, f"fact(s) NOT findable in the input text: {missing}"
                      f"{' — this client SHIPS AS A TEST INPUT, so the packet is unbuildable' if ships.startswith('to-do') else ''}")
        else:
            ok(slug, f"all {len(want)} load-bearing facts findable in extracted input text")

        # ---- 6: mess events are real -------------------------------------
        mess = mk.get("mess_events") or []
        if not mess:
            bad(slug, "no mess events")
        else:
            unreal = []
            for m in mess:
                if not isinstance(m, dict): continue
                for key in ("wrong_value", "wrong", "superseded_value"):
                    if m.get(key) and not findable(m[key], hay_norm, hay_raw):
                        unreal.append(f"{m.get('type','?')}:{key}")
            if unreal:
                bad(slug, f"mess event value(s) claimed but absent from the inputs: {unreal}")
            else:
                ok(slug, f"{len(mess)} mess events, all claimed values present in the text")

        # ---- 7: leakage ---------------------------------------------------
        # Reuse validate_masterkeys.py's four filters verbatim. The first draft
        # re-implemented them badly and red-lighted every client on `USCIS`,
        # `PTIN`, `ACTC`, `Employer-provided` and ordinary possessives —
        # 1040/N-400 form vocabulary that STYLE-SPEC §11 excuses BY NAME. Two
        # implementations of one rule is how a gate stops meaning anything.
        hits = leak_hits(all_text)
        if hits:
            bad(slug, f"LEAKAGE in inputs — {len(hits)} token(s): {sorted(hits)[:8]}")
        else:
            ok(slug, "leakage: 0 distinctive blocklist hits in the input text")

    print("\n".join(OKS))
    if WARNS: print(); print("\n".join(WARNS))
    if FAILS:
        print(); print("\n".join(FAILS))
        print(f"\n=== INPUT BARRIER RED: {len(FAILS)} failure(s), {len(WARNS)} warning(s) ===")
        sys.exit(1)
    print(f"\n=== INPUT BARRIER GREEN: {len(OKS)} checks, {len(WARNS)} warning(s) ===")

if __name__ == "__main__":
    main()

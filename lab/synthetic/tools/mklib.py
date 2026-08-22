#!/usr/bin/env python3
"""mklib.py — the shared spine every Phase 3 renderer imports.

Contract: tools/RENDER-CONTRACT.md. Nothing here decides house style; it
enforces the one spelling of a thing that STYLE-SPEC / SPEC-DELTA already
decided, so that five concurrently-built renderers cannot drift.

Authority order: STYLE-SPEC §16 > spec/SPEC-DELTA.md > tools/n400-part-map.md
> templates/document-catalog.yaml > the rest of STYLE-SPEC.
"""
from __future__ import annotations

import datetime
import os
import re
import shutil
import subprocess
import tempfile

import yaml

# --------------------------------------------------------------- paths
TOOLS = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(TOOLS)                     # lab/synthetic
CLIENTS = os.path.join(ROOT, "clients")
BLANKS = os.path.join(ROOT, "blanks")
TEMPLATES = os.path.join(ROOT, "templates")
SPEC = os.path.join(ROOT, "spec")

N400_BLANK = os.path.join(BLANKS, "n-400.pdf")

# ------------------------------------------------- geometry / typography
# STYLE-SPEC §3. US Letter portrait, 1-inch margins, Times New Roman 12,
# line spacing 1.15.
PAGE_W, PAGE_H = 612.0, 792.0
MARGIN = 72.0
BODY_FONT = "Times-Roman"
BODY_FONT_BOLD = "Times-Bold"
BODY_PT = 12.0
LEADING = 13.8                       # 12 pt x 1.15
DOCX_FONT = "Times New Roman"
DOCX_LINE_SPACING = 1.15

# Divider and tab-cover point sizes are NOT the same. §16 ruling 3 lowered the
# dividers to 12 pt and its consequences note is explicit that the tab cover
# pages are untouched (SPEC-DELTA §A, "Explicitly NOT changed").
DIVIDER_PT = 12.0
TAB_COVER_PT = 24.0

TAB_FOLDER = {"A": "Tab A (Content + Cover)", "B": "Tab B (Biographical Info)"}
MERGED_NAME = "N-400 Packet.pdf"

# ----------------------------------------------------------- determinism
# RENDER-CONTRACT §6. One creation date for the whole build. Never now().
FIXED_PDF_DATE = "D:20260101000000Z"
_FIXED_ID = b"\x00" * 16
SOURCE_DATE_EPOCH = "1767225600"     # 2026-01-01T00:00:00Z, for soffice
FIXED_DOCX_DATETIME = datetime.datetime(2026, 1, 1, 0, 0, 0)
FIXED_ZIP_DATE = (2026, 1, 1, 0, 0, 0)


# =====================================================================
# loading
# =====================================================================
def load_masterkey(slug: str) -> dict:
    """The ONLY masterkey reader. Reads masterkey.norm.yaml and nothing else.

    SPEC-DELTA D-I: the six authored masterkey.yaml files are six different
    shapes. Everything from Phase 3 onward reads the normalised file.
    """
    p = os.path.join(CLIENTS, slug, "masterkey.norm.yaml")
    if not os.path.exists(p):
        raise FileNotFoundError(
            f"{p} missing — run normalize_masterkeys.py. "
            "Renderers never read the raw masterkey.yaml (SPEC-DELTA D-I).")
    with open(p) as fh:
        return yaml.safe_load(fh)


def all_slugs() -> list[str]:
    return sorted(d for d in os.listdir(CLIENTS)
                  if os.path.isdir(os.path.join(CLIENTS, d)))


def load_catalog() -> dict:
    with open(os.path.join(TEMPLATES, "document-catalog.yaml")) as fh:
        return yaml.safe_load(fh)


def load_template(name: str) -> dict:
    if not name.endswith(".yaml"):
        name += ".yaml"
    with open(os.path.join(TEMPLATES, name)) as fh:
        return yaml.safe_load(fh)


# =====================================================================
# the joined `doc` dict — RENDER-CONTRACT §3
# =====================================================================
_SLOT = re.compile(r"\{([A-Z_]+)\}")


def _resolve_slots(text: str, mk: dict) -> str:
    """Resolve the catalog's template slots. Currently only {TAX_YEAR}.

    Done here, once, so doc['file_stem'] and doc['divider_title'] are ALWAYS
    literal by the time a renderer sees them.
    """
    if not isinstance(text, str) or "{" not in text:
        return text
    subs = {"TAX_YEAR": str(((mk.get("documents") or {})
                             .get("tax_return") or {}).get("year", ""))}

    def rep(m):
        key = m.group(1)
        if key not in subs or not subs[key]:
            raise KeyError(f"catalog slot {{{key}}} has no value for this client")
        return subs[key]

    return _SLOT.sub(rep, text)


def doc_entries(mk: dict) -> list[dict]:
    """This client's DOCUMENTs, ordered by seq, catalog joined to masterkey.

    mk['exhibits'] is the authority for the document set. mk['exhibits_derived']
    is advisory and is ABSENT on tran_daniel — never read it.
    """
    cat = {d["id"]: d for d in load_catalog()["documents"]}
    out = []
    for e in (mk.get("exhibits") or []):
        if not isinstance(e, dict):
            raise TypeError(f"exhibits entry is {type(e).__name__}, expected dict "
                            "— masterkey is not normalised")
        doc_id = e.get("doc")
        if doc_id not in cat:
            raise KeyError(f"exhibits names {doc_id!r}, absent from document-catalog.yaml")
        c = cat[doc_id]
        out.append({
            "id": doc_id,
            "seq": int(e["seq"]),
            "tab": c["tab"],
            "trigger": e.get("trigger", c.get("trigger", "core")),
            "file_stem": _resolve_slots(c["file_stem"], mk),
            "divider_title": _resolve_slots(c["divider_title"], mk),
            "toc_line": c["toc_line"],
            "ships": list(c["ships"]),
            "authored_by": c["authored_by"],
            "renderer": c.get("renderer"),
            "notes": c.get("notes"),
            "why": e.get("why", ""),
        })
    out.sort(key=lambda d: d["seq"])
    seqs = [d["seq"] for d in out]
    if seqs != list(range(1, len(seqs) + 1)):
        raise ValueError(f"DOCUMENT numbers are not 1..n contiguous: {seqs}")
    return out


def doc_by_id(mk: dict, doc_id: str) -> dict:
    for d in doc_entries(mk):
        if d["id"] == doc_id:
            return d
    raise KeyError(f"{doc_id} is not a DOCUMENT for this client")


def pre_documents() -> list[dict]:
    """The three pre-documents: no DOCUMENT number, no divider."""
    return list(load_catalog()["pre_documents"])


# =====================================================================
# paths — the ONLY place a component file name is spelled
# =====================================================================
def ensure_outdir(outdir: str) -> str:
    outdir = os.path.abspath(outdir)
    os.makedirs(outdir, exist_ok=True)
    for tab in TAB_FOLDER:
        os.makedirs(os.path.join(outdir, TAB_FOLDER[tab]), exist_ok=True)
    return outdir


def tab_dir(outdir: str, tab: str) -> str:
    if tab not in TAB_FOLDER:
        raise ValueError(f"tab must be 'A' or 'B', got {tab!r}")
    p = os.path.join(os.path.abspath(outdir), TAB_FOLDER[tab])
    os.makedirs(p, exist_ok=True)
    return p


def component_path(outdir: str, doc: dict, ext: str) -> str:
    """<TAB>-<seq>. <file_stem>.<ext>   — the CONTENT file (STYLE-SPEC §2)."""
    ext = ext.lstrip(".")
    return os.path.join(tab_dir(outdir, doc["tab"]),
                        f'{doc["tab"]}-{doc["seq"]}. {doc["file_stem"]}.{ext}')


def divider_path(outdir: str, doc: dict) -> str:
    """<TAB>-<seq>.pdf — the BARE numbered pdf, which is the DIVIDER.

    render_docs.py only. RENDER-CONTRACT §0.4: no other renderer writes this.
    """
    return os.path.join(tab_dir(outdir, doc["tab"]),
                        f'{doc["tab"]}-{doc["seq"]}.pdf')


def merged_path(outdir: str) -> str:
    return os.path.join(os.path.abspath(outdir), MERGED_NAME)


# =====================================================================
# dates — exactly two house formats
# =====================================================================
_MONTHS = ("January", "February", "March", "April", "May", "June", "July",
           "August", "September", "October", "November", "December")


def as_date(v) -> datetime.date:
    if isinstance(v, datetime.datetime):
        return v.date()
    if isinstance(v, datetime.date):
        return v
    if isinstance(v, str):
        s = v.strip()
        for f in ("%Y-%m-%d", "%m/%d/%Y", "%B %d, %Y", "%Y/%m/%d"):
            try:
                return datetime.datetime.strptime(s, f).date()
            except ValueError:
                pass
    raise ValueError(f"cannot read {v!r} as a date")


def fmt_numeric(v) -> str:
    """MM/DD/YYYY — the form/cover-page format."""
    d = as_date(v)
    return f"{d.month:02d}/{d.day:02d}/{d.year:04d}"


def fmt_long(v) -> str:
    """Month D, YYYY — the cover-letter format. No leading zero on the day."""
    d = as_date(v)
    return f"{_MONTHS[d.month - 1]} {d.day}, {d.year}"


# =====================================================================
# determinism
# =====================================================================
def stamp_deterministic(target) -> None:
    """Fix /CreationDate, /ModDate and the file ID.

    target: a pypdf PdfWriter (stamped in place, before .write), or a path to
    an existing PDF (rewritten in place).
    """
    from pypdf import PdfWriter
    from pypdf.generic import (ArrayObject, ByteStringObject, NameObject,
                               TextStringObject)

    def _apply(w):
        from pypdf.generic import DictionaryObject
        info = {NameObject("/CreationDate"): TextStringObject(FIXED_PDF_DATE),
                NameObject("/ModDate"): TextStringObject(FIXED_PDF_DATE),
                NameObject("/Producer"): TextStringObject("mklib"),
                NameObject("/Creator"): TextStringObject("mklib")}
        di = w._info.get_object() if hasattr(w._info, "get_object") else w._info
        if not isinstance(di, DictionaryObject):
            di = DictionaryObject()
            w._info = w._add_object(di)
        for k, val in info.items():
            di[k] = val
        try:
            w._ID = ArrayObject([ByteStringObject(_FIXED_ID),
                                 ByteStringObject(_FIXED_ID)])
        except Exception:
            pass
        # LibreOffice writes an XMP packet carrying its own wall clock
        # (<xmp:CreateDate>), which /Info does NOT cover. Two identical
        # documents converted eight seconds apart differ in exactly those
        # bytes. The /Info dictionary above carries everything we need, so
        # drop the XMP stream rather than try to rewrite dates inside it.
        try:
            if "/Metadata" in w._root_object:
                del w._root_object[NameObject("/Metadata")]
        except Exception:
            pass

    if isinstance(target, str):
        w = PdfWriter(clone_from=target)
        _apply(w)
        tmp = target + ".det.tmp"
        with open(tmp, "wb") as fh:
            w.write(fh)
        os.replace(tmp, target)
        _freeze_xmp(target)
    else:
        _apply(target)


_XMP_DATE = re.compile(
    rb"(<(?:xmp|xap):(?:CreateDate|ModifyDate|MetadataDate)>)"
    rb"(\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2})"
    rb"(Z|[+-]\d{2}:\d{2})")
_FIXED_DATETIME = b"2026-01-01T00:00:00"

# Ghostscript's pdfwrite (merge_packet.py's flatten step) also mints a fresh
# xmpMM:DocumentID (a UUID1 — its time-based octets differ run to run even
# with every date above frozen) on every invocation. Neutralise it the same
# length-preserving way: keep dash positions, zero the hex digits.
_XMP_UUID = re.compile(
    rb"((?:DocumentID|InstanceID)=['\"]uuid:)([0-9a-fA-F-]+)(['\"])")


def _zero_uuid(hexdashes: bytes) -> bytes:
    return bytes((b if b == ord("-") else ord("0")) for b in hexdashes)


def _freeze_xmp(path: str) -> None:
    """Replace the wall clock (and Ghostscript's fresh DocumentID) inside an
    embedded XMP packet.

    The XMP stream is a SECOND place a timestamp lives, independent of /Info,
    and pypdf keeps the object alive even when /Metadata is unlinked from the
    catalog. LibreOffice writes a 'Z'-suffixed UTC timestamp; Ghostscript's
    pdfwrite writes one with a local zone-offset suffix ("...-04:00") instead
    of 'Z' — same field, different tail. The date/time portion is replaced
    with a fixed value of identical length; the original 'Z' or offset tail
    is preserved verbatim. The DocumentID substitution is the same
    length-preserving trick, hex digit for hex digit. Both keep every byte
    offset and the xref table valid regardless of which tool produced the
    packet.
    """
    with open(path, "rb") as fh:
        data = fh.read()
    if b"CreateDate" not in data and b"DocumentID" not in data:
        return
    patched = _XMP_DATE.sub(
        lambda m: m.group(1) + _FIXED_DATETIME + m.group(3), data)
    patched = _XMP_UUID.sub(
        lambda m: m.group(1) + _zero_uuid(m.group(2)) + m.group(3), patched)
    if patched != data and len(patched) == len(data):
        with open(path, "wb") as fh:
            fh.write(patched)


# =====================================================================
# reportlab
# =====================================================================
def new_canvas(path: str):
    """US Letter portrait, Times-Roman 12, deterministic.

    reportlab's invariant=1 suppresses its own timestamp and document ID.
    """
    from reportlab.pdfgen import canvas
    os.makedirs(os.path.dirname(os.path.abspath(path)) or ".", exist_ok=True)
    c = canvas.Canvas(path, pagesize=(PAGE_W, PAGE_H), invariant=1,
                      pageCompression=1)
    c.setFont(BODY_FONT, BODY_PT)
    return c


def frame_text(c, lines, align="left", top=None, size=BODY_PT,
               leading=None, font=None):
    """Lay a list of strings inside the 1-inch frame.

    lines: str, or (str, {"bold":bool}) tuples. Returns the next free y.
    """
    leading = leading if leading is not None else size * 1.15
    y = (PAGE_H - MARGIN - size) if top is None else top
    for item in lines:
        opts = {}
        if isinstance(item, tuple):
            item, opts = item
        f = font or (BODY_FONT_BOLD if opts.get("bold") else BODY_FONT)
        sz = opts.get("size", size)
        c.setFont(f, sz)
        text = "" if item is None else str(item)
        if text:
            if align == "center":
                c.drawCentredString(PAGE_W / 2.0, y, text)
            elif align == "right":
                c.drawRightString(PAGE_W - MARGIN, y, text)
            else:
                c.drawString(MARGIN + opts.get("indent", 0.0), y, text)
            if opts.get("underline"):
                w = c.stringWidth(text, f, sz)
                x0 = (PAGE_W - w) / 2.0 if align == "center" else MARGIN
                c.line(x0, y - 2, x0 + w, y - 2)
        y -= leading
    return y


# =====================================================================
# docx  (authored_by: firm only)
# =====================================================================
def new_docx():
    """Letter, 1" margins, Times New Roman 12, 1.15 spacing, no header/footer."""
    import docx
    from docx.shared import Inches, Pt

    d = docx.Document()
    for s in d.sections:
        s.page_width, s.page_height = Inches(8.5), Inches(11)
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)
    st = d.styles["Normal"]
    st.font.name = DOCX_FONT
    st.font.size = Pt(BODY_PT)
    rpr = st.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rpr.set(_w(attr), DOCX_FONT)
    st.paragraph_format.line_spacing = DOCX_LINE_SPACING
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.space_before = Pt(0)
    # Determinism (RENDER-CONTRACT §6): python-docx otherwise stamps
    # docProps/core.xml with datetime.now(), which changes the .docx AND every
    # pdf soffice derives from it.
    cp = d.core_properties
    cp.created = FIXED_DOCX_DATETIME
    cp.modified = FIXED_DOCX_DATETIME
    cp.last_modified_by = ""
    cp.revision = 1
    cp.author = ""
    return d


def save_docx(document, path: str) -> str:
    """Save a docx and freeze its zip entry timestamps.

    python-docx writes each zip member with the current clock, so two identical
    documents produce two different files. Rewrite the archive with one fixed
    date, member order preserved. Always save through here.
    """
    import zipfile
    path = os.path.abspath(path)
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    document.save(path)
    tmp = path + ".det.tmp"
    with zipfile.ZipFile(path) as zin:
        infos = zin.infolist()
        data = {i.filename: zin.read(i.filename) for i in infos}
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for i in infos:
            zi = zipfile.ZipInfo(i.filename, date_time=FIXED_ZIP_DATE)
            zi.compress_type = i.compress_type
            zi.external_attr = i.external_attr
            zi.create_system = 3
            zout.writestr(zi, data[i.filename])
    os.replace(tmp, path)
    return path


def _w(tag: str) -> str:
    return ("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            + tag)


_ALIGN = {"left": 0, "center": 1, "right": 2, "both": 3}


def add_para(document, text="", align="left", bold=False, underline=False,
             indent_in=0.0, size_pt=None, spacing_after=None, tab_stops_in=()):
    """A plain paragraph. No text boxes, no tables, no images (STYLE-SPEC §3)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH(_ALIGN[align])
    pf = p.paragraph_format
    pf.line_spacing = DOCX_LINE_SPACING
    if indent_in:
        pf.left_indent = Inches(indent_in)
    if spacing_after is not None:
        pf.space_after = Pt(spacing_after)
    for stop in tab_stops_in:
        pf.tab_stops.add_tab_stop(Inches(stop))
    if text:
        r = p.add_run(text)
        r.bold = bool(bold)
        r.underline = bool(underline)
        r.font.name = DOCX_FONT
        if size_pt:
            r.font.size = Pt(size_pt)
        rf = r._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rf.set(_w(attr), DOCX_FONT)
    return p


def add_blank(document, size_pt=None):
    """The corpus's only block separator: an EMPTY PARAGRAPH (STYLE-SPEC §3)."""
    return add_para(document, "", size_pt=size_pt)


# Per-process, NOT shared: LibreOffice refuses to start a second instance
# against a profile another instance holds, and Phase 3 runs five renderers
# concurrently. The profile affects nothing in the output, so a private one
# costs no determinism.
_SOFFICE_PROFILE = os.path.join(tempfile.gettempdir(),
                                f"mklib-soffice-profile-{os.getpid()}")


def docx_to_pdf(docx_path: str, outdir: str | None = None) -> str:
    """soffice --headless --convert-to pdf, then normalised deterministic."""
    docx_path = os.path.abspath(docx_path)
    outdir = os.path.abspath(outdir or os.path.dirname(docx_path))
    os.makedirs(outdir, exist_ok=True)
    env = dict(os.environ, SOURCE_DATE_EPOCH=SOURCE_DATE_EPOCH, TZ="UTC")
    # Convert into a scratch dir: soffice derives the output name from the
    # input stem, and our stems contain dots ("A-1. Table of Contents").
    with tempfile.TemporaryDirectory() as scratch:
        r = subprocess.run(
            ["soffice", "--headless", "--norestore",
             f"-env:UserInstallation=file://{_SOFFICE_PROFILE}",
             "--convert-to", "pdf", "--outdir", scratch, docx_path],
            capture_output=True, text=True, env=env, timeout=300)
        produced = [f for f in os.listdir(scratch) if f.lower().endswith(".pdf")]
        if not produced:
            # One retry: a concurrent soffice can still lose a race on shared
            # config even with a private profile.
            r = subprocess.run(
                ["soffice", "--headless", "--norestore",
                 f"-env:UserInstallation=file://{_SOFFICE_PROFILE}",
                 "--convert-to", "pdf", "--outdir", scratch, docx_path],
                capture_output=True, text=True, env=env, timeout=300)
            produced = [f for f in os.listdir(scratch) if f.lower().endswith(".pdf")]
        if not produced:
            raise RuntimeError(
                f"soffice produced no pdf for {docx_path}\n{r.stdout}\n{r.stderr}")
        dest = os.path.join(outdir,
                            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        shutil.move(os.path.join(scratch, produced[0]), dest)
    stamp_deterministic(dest)
    return dest


# =====================================================================
# pdf utilities
# =====================================================================
def pdf_pagecount(path: str) -> int:
    from pypdf import PdfReader
    return len(PdfReader(path).pages)


def pdf_text(path: str, first: int | None = None, last: int | None = None) -> str:
    cmd = ["pdftotext", "-layout"]
    if first:
        cmd += ["-f", str(first)]
    if last:
        cmd += ["-l", str(last)]
    cmd += [path, "-"]
    return subprocess.run(cmd, capture_output=True, text=True).stdout


# =====================================================================
# AcroForm
# =====================================================================
def field_names(path: str) -> set:
    """Fully-qualified AcroForm field names, exactly as pypdf keys them.

    Authoritative — prefer this to parsing n400-field-dump.tsv, whose tooltip
    column carries embedded newlines that break a naive line splitter.
    """
    from pypdf import PdfReader
    return set((PdfReader(path).get_fields() or {}).keys())


def field_values(path: str) -> dict:
    from pypdf import PdfReader
    out = {}
    for k, v in (PdfReader(path).get_fields() or {}).items():
        val = v.get("/V")
        out[k] = "" if val is None else str(val)
    return out


def btn_on_states(path: str) -> dict:
    """field name -> its single 'on' appearance state, e.g. '/Y', '/N', '/APT'.

    A /Btn on these blanks is NOT a Yes/No checkbox. Each widget's /AP /N
    dictionary holds exactly one key and that key is the only value that checks
    the box; anything else stores as /Off and the box prints blank. The Yes and
    No boxes are separate FIELDS, and the index in the name is not the box
    letter (Part1_Eligibility[2] is box A).
    """
    from pypdf import PdfReader
    out = {}
    for page in PdfReader(path).pages:
        for annot in (page.get("/Annots") or []):
            o = annot.get_object()
            node, ft = o, None
            while node is not None:
                if "/FT" in node:
                    ft = node["/FT"]
                    break
                node = node.get("/Parent")
            if ft != "/Btn":
                continue
            name, node = [], o
            while node is not None:
                if "/T" in node:
                    name.insert(0, str(node["/T"]))
                node = node.get("/Parent")
            ap = o.get("/AP", {})
            n = ap.get("/N", {}) if hasattr(ap, "get") else {}
            states = [str(k) for k in n.keys()] if hasattr(n, "keys") else []
            on = [s for s in states if s != "/Off"]
            if on:
                out[".".join(name)] = on[0]
    return out


def fill_acroform(src: str, dst: str, values: dict) -> None:
    """Fill, delete /XFA, set NeedAppearances, stamp the fixed date, write dst.

    STYLE-SPEC §8: born-digital. Values are set on every page so that a field
    whose widget lives on a page other than the one guessed still takes.
    """
    from pypdf import PdfWriter
    from pypdf.generic import BooleanObject, NameObject

    w = PdfWriter(clone_from=src)
    acro = w._root_object.get("/AcroForm")
    if acro is not None and hasattr(acro, "get_object"):
        acro = acro.get_object()
    if acro is not None:
        acro[NameObject("/NeedAppearances")] = BooleanObject(True)
        if "/XFA" in acro:
            del acro[NameObject("/XFA")]
    for page in w.pages:
        w.update_page_form_field_values(page, values, auto_regenerate=False)
    n = _build_text_appearances(w)
    # Appearances are ours now; a regenerator would undo them.
    if acro is not None and n:
        acro[NameObject("/NeedAppearances")] = BooleanObject(False)
    stamp_deterministic(w)
    os.makedirs(os.path.dirname(os.path.abspath(dst)) or ".", exist_ok=True)
    with open(dst, "wb") as fh:
        w.write(fh)


# Helvetica widths (1000 em units) for the printable ASCII range, so a value's
# width can be measured rather than guessed. Estimating with a flat 0.5 em was
# the first attempt and it under-shrank every wide cell.
_HELV_W = {
    ' ': 278, '!': 278, '"': 355, '#': 556, '$': 556, '%': 889, '&': 667,
    "'": 191, '(': 333, ')': 333, '*': 389, '+': 584, ',': 278, '-': 333,
    '.': 278, '/': 278, ':': 278, ';': 278, '<': 584, '=': 584, '>': 584,
    '?': 556, '@': 1015, '[': 278, '\\': 278, ']': 278, '^': 469, '_': 556,
    '`': 333, '{': 334, '|': 260, '}': 334, '~': 584,
}
def _seed_helv_widths():
    """Populated in a function so its loop variables cannot leak into module
    scope. The first version used `_w` as the loop name and silently rebound
    mklib's own `_w()` XML-namespace helper to an int, which broke every docx
    page with 'int object is not callable'."""
    for ch in '0123456789':
        _HELV_W[ch] = 556
    for ch, wid in zip('ABCDEFGHIJKLMNOPQRSTUVWXYZ',
                       (667, 667, 722, 722, 667, 611, 778, 722, 278, 500, 667,
                        556, 833, 722, 778, 667, 778, 722, 667, 611, 722, 667,
                        944, 667, 667, 611)):
        _HELV_W[ch] = wid
    for ch, wid in zip('abcdefghijklmnopqrstuvwxyz',
                       (556, 556, 500, 556, 556, 278, 556, 556, 222, 222, 500,
                        222, 833, 556, 556, 556, 556, 333, 500, 278, 556, 500,
                        722, 500, 500, 500)):
        _HELV_W[ch] = wid


_seed_helv_widths()


def _helv_width(text: str, size: float) -> float:
    return sum(_HELV_W.get(ch, 556) for ch in text) * size / 1000.0


def _build_text_appearances(w) -> int:
    """Give every filled text field a real /AP appearance stream.

    WHY THIS EXISTS. `NeedAppearances = true` tells a *viewer* to generate the
    appearance of each field. Ghostscript, doing the flatten, generates them at
    its own fixed size and clips whatever does not fit. The N-400's Part 4 and
    Part 7 tables have narrow cells, so legitimate values were printed as
    "Harborline Struc", "United St", "ructural Engine" — correct in the field
    data and correct under `pdftotext`, wrong on the page a human reads. Every
    text-based gate passed it; a Phase 5 reviewer looking at a rendered image
    caught it. Setting a per-field /DA did not help, because with
    NeedAppearances set gs never consults it.

    So we stop asking and draw them ourselves: measure the string in Helvetica,
    pick the largest size from 9pt down to 4pt that fits the widget, and write
    the appearance stream. NeedAppearances is then turned OFF, because a
    generator that regenerates appearances would undo this.

    Returns the number of widgets given an appearance.
    """
    from pypdf.generic import (ArrayObject, DecodedStreamObject, DictionaryObject,
                               FloatObject, NameObject, NumberObject)
    BASE, FLOOR, PAD = 9.0, 4.0, 2.0
    helv = DictionaryObject({
        NameObject("/Type"): NameObject("/Font"),
        NameObject("/Subtype"): NameObject("/Type1"),
        NameObject("/BaseFont"): NameObject("/Helvetica"),
        NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
    })
    helv_ref = w._add_object(helv)
    made = 0
    for page in w.pages:
        for a in (page.get("/Annots") or []):
            try:
                o = a.get_object()
            except Exception:
                continue
            # Never touch the 2D barcode fields. They carry their own graphic
            # appearance; replacing it printed the literal string
            # "N-400|01/20/25|5" where the PDF417 block belongs.
            nm, cur = [], o
            while cur is not None:
                t = cur.get("/T")
                if t: nm.append(str(t))
                nxt = cur.get("/Parent")
                cur = nxt.get_object() if nxt is not None else None
            fullname = ".".join(reversed(nm))
            if "BarCode" in fullname or "Barcode" in fullname:
                continue
            node, val, ft, ff = o, None, None, 0
            while node is not None:
                if val is None and node.get("/V") is not None:
                    val = node.get("/V")
                if ft is None and node.get("/FT") is not None:
                    ft = node.get("/FT")
                if node.get("/Ff") is not None:
                    try: ff = int(node.get("/Ff"))
                    except Exception: pass
                nxt = node.get("/Parent")
                node = nxt.get_object() if nxt is not None else None
            if ft != "/Tx" or not isinstance(val, str) or not val.strip():
                continue
            if ff & (1 << 24):        # comb field — the form spaces it itself
                continue
            rect = o.get("/Rect")
            if not rect or len(rect) != 4:
                continue
            try:
                x0, y0, x1, y1 = (float(rect[0]), float(rect[1]),
                                  float(rect[2]), float(rect[3]))
            except Exception:
                continue
            bw, bh = abs(x1 - x0), abs(y1 - y0)
            if bw <= 2 * PAD or bh <= 2:
                continue
            text = val.strip()
            size = min(BASE, bh * 0.68)
            avail = bw - 2 * PAD
            if _helv_width(text, size) > avail:
                size = max(FLOOR, avail / max(_helv_width(text, 1.0), 1e-6))
                size = min(size, bh * 0.68)
            ty = (bh - size) / 2.0 + size * 0.22
            esc = (text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)"))
            stream = (f"/Tx BMC q BT /Helv {size:.2f} Tf 0 g "
                      f"{PAD:.2f} {ty:.2f} Td ({esc}) Tj ET Q EMC")
            ap = DecodedStreamObject()
            ap.set_data(stream.encode("latin-1", "replace"))
            ap[NameObject("/Type")] = NameObject("/XObject")
            ap[NameObject("/Subtype")] = NameObject("/Form")
            ap[NameObject("/FormType")] = NumberObject(1)
            ap[NameObject("/BBox")] = ArrayObject(
                [FloatObject(0), FloatObject(0), FloatObject(bw), FloatObject(bh)])
            ap[NameObject("/Resources")] = DictionaryObject({
                NameObject("/Font"): DictionaryObject({NameObject("/Helv"): helv_ref}),
            })
            o[NameObject("/AP")] = DictionaryObject({NameObject("/N"): w._add_object(ap)})
            made += 1
    return made


# =====================================================================
# MRZ — ICAO 9303
# =====================================================================
def mrz_check_digit(s: str) -> int:
    """Weights 7,3,1. Digits as themselves, A=10..Z=35, '<'=0."""
    w, tot = (7, 3, 1), 0
    for i, ch in enumerate(s):
        if ch.isdigit():
            v = int(ch)
        elif ch.isalpha():
            v = ord(ch.upper()) - 55
        elif ch == "<":
            v = 0
        else:
            raise ValueError(f"illegal MRZ character {ch!r}")
        tot += v * w[i % 3]
    return tot % 10


def _mrz_date(v) -> str:
    d = as_date(v)
    return f"{d.year % 100:02d}{d.month:02d}{d.day:02d}"


def mrz_lines(surname, given_names, doc_number, nationality, dob, sex,
              expiry, personal_number="") -> tuple:
    """Two 44-char TD3 lines with all four check digits computed.

    Use this to ASSERT the masterkey's stored MRZ recomputes. Print the
    masterkey's lines; never invent a different pair.
    """
    def clean(s):
        return re.sub(r"[^A-Z]", "<", str(s).upper())

    name = f"{clean(surname)}<<{'<'.join(clean(g) for g in str(given_names).split())}"
    line1 = ("P<" + clean(nationality)[:3] + name).ljust(44, "<")[:44]

    num = str(doc_number).upper().ljust(9, "<")[:9]
    pn = re.sub(r"[^A-Z0-9]", "<", str(personal_number).upper()).ljust(14, "<")[:14]
    body = (num + str(mrz_check_digit(num))
            + clean(nationality)[:3].ljust(3, "<")
            + _mrz_date(dob) + str(mrz_check_digit(_mrz_date(dob)))
            + str(sex).upper()[:1]
            + _mrz_date(expiry) + str(mrz_check_digit(_mrz_date(expiry)))
            + pn + str(mrz_check_digit(pn)))
    composite = body[0:10] + body[13:20] + body[21:43]
    line2 = body + str(mrz_check_digit(composite))
    return line1, line2


def mrz_verify(line1: str, line2: str) -> list:
    """Return a list of human-readable problems; empty means all four agree."""
    problems = []
    for nm, ln in (("line1", line1), ("line2", line2)):
        if len(ln) != 44:
            problems.append(f"MRZ {nm} is {len(ln)} chars, must be 44")
    if problems:
        return problems
    for nm, data, got in (("passport-number", line2[0:9], line2[9]),
                          ("dob", line2[13:19], line2[19]),
                          ("expiry", line2[21:27], line2[27])):
        exp = mrz_check_digit(data)
        if not got.isdigit() or int(got) != exp:
            problems.append(f"MRZ {nm} check digit {got!r}, computed {exp}")
    comp = line2[0:10] + line2[13:20] + line2[21:43]
    exp = mrz_check_digit(comp)
    if not line2[43].isdigit() or int(line2[43]) != exp:
        problems.append(f"MRZ composite check digit {line2[43]!r}, computed {exp}")
    return problems


# =====================================================================
# lockbox — STYLE-SPEC §7, as amended by §16 ruling 8
# =====================================================================
# §16 ruling 8 OVERTURNED taste call 8: the block opens with jacobs' longer
# form, "U.S. Department of Homeland Security". templates/cover-letter.docx.yaml
# still carried zhu's shorter form; §16 wins (see TOOLSMITH-NOTES in the Phase 3
# report). This is the only place either string is spelled.
DHS_LINES = ["U.S. Department of Homeland Security",
             "United States Citizenship and Immigration Services"]

LOCKBOX_STATES = {
    "Elgin": "CT DE DC FL GA ME MD MA NH NJ NY NC PA RI SC VT VA WV".split(),
    "Phoenix": ("AL AK AZ CA CO HI ID KS KY MN MS MT NE NV NM ND OR SD TN UT "
                "WA WY AS GU MP PR VI AA AE AP").split(),
    "Dallas": "AR LA OK TX".split(),
    "Chicago": "IL IN IA MI MO OH WI".split(),
}

_LOCKBOX_ADDR = {
    ("Elgin", "usps"): ["USCIS Elgin Lockbox", "Attn: N-400",
                        "P.O. Box 4060", "Carol Stream, IL 60197-4060"],
    ("Elgin", "courier"): ["USCIS Elgin Lockbox", "Attn: N-400 (Box 4060)",
                           "2500 Westfield Drive", "Elgin, IL 60124-7836"],
    ("Phoenix", "usps"): ["USCIS Phoenix Lockbox", "Attn: N-400",
                          "P.O. Box 21251", "Phoenix, AZ 85036-1251"],
    ("Phoenix", "courier"): ["USCIS Phoenix Lockbox", "Attn: N-400 (Box 21251)",
                             "2108 E. Elliot Rd.", "Tempe, AZ 85284-1806"],
    ("Dallas", "usps"): ["USCIS Dallas Lockbox", "Attn: N-400",
                         "P.O. Box 660060", "Dallas, TX 75266-0060"],
    ("Dallas", "courier"): ["USCIS Dallas Lockbox", "Attn: N-400 (Box 660060)",
                            "2501 S State Hwy 121 Business", "Suite 400",
                            "Lewisville, TX 75067-8003"],
    ("Chicago", "usps"): ["USCIS Chicago Lockbox", "Attn: N-400",
                          "P.O. Box 4380", "Chicago, IL 60680-4380"],
    ("Chicago", "courier"): ["USCIS Chicago Lockbox", "Attn: N-400 (Box 4380)",
                             "131 S. Dearborn, 3rd Floor", "Chicago, IL 60603-5517"],
}

CARRIERS = ("U.S. POSTAL SERVICE (USPS)", "FEDERAL EXPRESS")


def lockbox_for_state(state: str) -> str:
    st = str(state).strip().upper()
    for box, states in LOCKBOX_STATES.items():
        if st in states:
            return box
    raise ValueError(f"no lockbox for state {state!r} (STYLE-SPEC §7)")


def lockbox_block(state: str, carrier: str) -> list:
    """The FULL address block, agency lines first. lockbox = f(state, carrier)."""
    if carrier not in CARRIERS:
        raise ValueError(f"carrier must be one of {CARRIERS}, got {carrier!r}")
    kind = "usps" if "POSTAL" in carrier.upper() else "courier"
    return list(DHS_LINES) + list(_LOCKBOX_ADDR[(lockbox_for_state(state), kind)])


def residence_state(mk: dict) -> str:
    for a in (mk.get("addresses") or []):
        if isinstance(a, dict) and a.get("present"):
            return a["state"]
    raise ValueError("no present address in masterkey")


# =====================================================================
# §5.1 — the eligibility clause and its citation
# =====================================================================
CITATIONS = {"319a": "8 U.S.C. § 1430(a); 8 C.F.R. § 319.1",
             "316a": "INA § 316(a); 8 C.F.R. § 316.2"}


def basis_key(mk: dict) -> str:
    b = str(mk["immigration"]["basis"]).replace("(", "").replace(")", "").lower()
    if "319" in b:
        return "319a"
    if "316" in b:
        return "316a"
    raise ValueError(f"unknown basis {b!r} — only 316(a) and 319(a) exist")


def eligibility_clause(mk: dict) -> str:
    """STYLE-SPEC §5.1. Slotted, no trailing period (the sentence supplies it)."""
    ident = mk["identity"]
    if basis_key(mk) == "319a":
        sp = mk["family"]["spouse"]
        full = sp.get("full_name") or sp.get("full_legal_name") or " ".join(
            x for x in (sp.get("given_name"), sp.get("middle_name"),
                        sp.get("family_name")) if x)
        return ("having three (3) years of permanent residency combined with "
                f"continuous residence with {ident['pronoun_possessive']} U.S. "
                f"citizen spouse, {sp['honorific']} {full}")
    return ("as it has been more than five (5) years since "
            f"{ident['pronoun_subject']} became a permanent resident on "
            f"{fmt_long(mk['immigration']['lpr_date'])}")


def citation(mk: dict) -> str:
    return CITATIONS[basis_key(mk)]


def classification_basis(mk: dict) -> str:
    """The applicant cover page's one-line basis. §16 r14: these two, nothing else."""
    return "INA 319(a)" if basis_key(mk) == "319a" else "INA 316(a)"


def full_legal_name(mk: dict) -> str:
    i = mk["identity"]
    return i.get("full_legal_name") or " ".join(
        x for x in (i.get("given_name"), i.get("middle_name"),
                    i.get("family_name")) if x)


# =====================================================================
# no-firm-identity guard (§16 r7) — used by verify_client.py and available
# to any renderer that wants to self-check.
# =====================================================================
FIRM_IDENTITY_TERMS = [
    # There is no firm name to look for, because none was ever invented
    # (§16 r7 killed §15.1). So we look for the SHAPES a preparer line takes.
    r"\bPrepared\s+by\b", r"\bAttorney\s+for\b", r"\bLaw\s+Offices?\b",
    r"\bLLP\b", r"\bEsq\.", r"\bAttorney\s+at\s+Law\b",
    r"\bAccredited\s+Representative\b", r"\bG-28\b",
    r"\bBy:\s*\S", r"\bOur\s+File\b", r"\bEncls?\.",
]


def firm_identity_hits(text: str) -> list:
    """Return (pattern, matched-text) for every firm-identity shape in text."""
    hits = []
    for pat in FIRM_IDENTITY_TERMS:
        for m in re.finditer(pat, text, re.I):
            hits.append((pat, m.group(0)))
    return hits


if __name__ == "__main__":
    import sys
    slug = sys.argv[1] if len(sys.argv) > 1 else "almeida_paulo"
    mk = load_masterkey(slug)
    print(f"{slug}: {len(doc_entries(mk))} documents, basis {basis_key(mk)}")
    for d in doc_entries(mk):
        print(f"  {d['tab']}-{d['seq']:<2} {d['trigger']:<5} "
              f"{d['file_stem']:<45} | {d['divider_title']}")
    print("  lockbox:", lockbox_block(residence_state(mk), mk["matter"]["carrier"]))
    print("  clause :", eligibility_clause(mk))

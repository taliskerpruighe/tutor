"""
lib/render.py -- FROZEN after Stage A. Do not edit from Stage B.

Rendering *mechanism*, not layout. Each company's Stage B worker brings its
own HTML/CSS templates, its own bank-statement look, its own filenames.
This module only supplies the machinery underneath:

  - render_html_to_pdf()   HTML+CSS -> text PDF, via weasyprint
  - scanify()              text PDF -> plausible image-only scanned PDF
  - photograph_receipt()   flat image -> photographed-looking JPG
  - handwritten_note_image()  short text -> handwritten-look JPG
  - concat_pdfs()          merge several PDFs into one multi-doc PDF
  - render_docx()          simple DOCX writer (title + paragraphs + table)
  - render_xlsx()          simple XLSX writer (dict of sheet -> rows)
  - write_csv()            CSV writer

Every function that makes a random choice takes a mandatory `seed: int`
and builds its own local `random.Random(seed)` -- never draws from the
global `random` module -- so a given seed always produces the same output
shape (not necessarily byte-identical: weasyprint and JPEG embed
timestamps/metadata).
"""

from __future__ import annotations

import csv
import io
import math
import os
import random
import subprocess
import tempfile

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as pdfcanvas


# ---------------------------------------------------------------------------
# HTML -> PDF
# ---------------------------------------------------------------------------

_WEASYPRINT_BIN = os.environ.get("WEASYPRINT_BIN", "weasyprint")


def render_html_to_pdf(html: str, out_path: str, css: str | None = None) -> None:
    """Render an HTML string (with optional CSS string) to a text PDF.

    Shells out to the `weasyprint` CLI rather than importing the `weasyprint`
    package: on this environment weasyprint is installed via pipx into its
    own isolated venv and is not importable from the system interpreter.
    """
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp_dir:
        html_path = os.path.join(tmp_dir, "in.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)
        cmd = [_WEASYPRINT_BIN]
        if css:
            css_path = os.path.join(tmp_dir, "style.css")
            with open(css_path, "w", encoding="utf-8") as f:
                f.write(css)
            cmd += ["-s", css_path]
        cmd += [html_path, out_path]
        subprocess.run(cmd, check=True, capture_output=True)


# ---------------------------------------------------------------------------
# scanify: text PDF -> plausible image-only scan
# ---------------------------------------------------------------------------

def _pdf_to_gray_pages(pdf_path: str, dpi: int, tmp_dir: str) -> list[str]:
    """Rasterise every page of pdf_path to grayscale PNGs via ghostscript.
    Returns list of PNG paths in page order. Ghostscript is used instead of
    ImageMagick/`convert` because convert's PDF delegate is commonly blocked
    by policy.xml in hardened installs."""
    out_pattern = os.path.join(tmp_dir, "page-%04d.png")
    subprocess.run(
        [
            "gs", "-dBATCH", "-dNOPAUSE", "-q",
            "-sDEVICE=pnggray",
            f"-r{dpi}",
            f"-sOutputFile={out_pattern}",
            pdf_path,
        ],
        check=True,
    )
    pages = sorted(
        os.path.join(tmp_dir, fn) for fn in os.listdir(tmp_dir) if fn.startswith("page-")
    )
    return pages


def scanify(
    src_pdf_path: str,
    out_pdf_path: str,
    seed: int,
    dpi: int = 200,
    jpeg_quality: int = 75,
    max_rotation_deg: float = 0.7,
) -> None:
    """Turn a text PDF into a plausible image-only scan: rasterise at `dpi`
    grayscale, JPEG-compress at `jpeg_quality`, apply a slight rotation, mild
    noise and contrast loss, then re-embed each page as a full-page image in
    a new image-only PDF. Must remain readable by tesseract: dates,
    counterparty and totals recoverable. Deterministic given `seed`.
    """
    rng = random.Random(seed)
    os.makedirs(os.path.dirname(out_pdf_path) or ".", exist_ok=True)

    with tempfile.TemporaryDirectory() as tmp_dir:
        page_pngs = _pdf_to_gray_pages(src_pdf_path, dpi, tmp_dir)
        writer = PdfWriter()

        for page_png in page_pngs:
            img = Image.open(page_png).convert("L")

            # Mild contrast loss: compress the tonal range toward mid-gray.
            lo, hi = 40, 225
            img = img.point(lambda p, lo=lo, hi=hi: lo + (p / 255.0) * (hi - lo))

            # Mild gaussian-ish noise (deterministic via local rng).
            #
            # NOTE on the size/quality trade-off: per-pixel iid noise (step=1)
            # blows the per-page JPEG size well past the 500KB/page raster
            # cap at 200 DPI/quality 75 on a text-dense page (measured
            # ~550KB). A coarser noise grid (step=6, small amplitude) reads
            # as the same "scanner grain" at normal viewing/OCR scale while
            # costing far less JPEG entropy (measured ~230-260KB on a dense
            # table page). Do not drop step below ~5 without re-measuring
            # per-page bytes on a dense fixture.
            noisy = img.copy()
            pixels = noisy.load()
            w, h = noisy.size
            noise_amount = 6
            step = 6
            for y in range(0, h, step):
                for x in range(0, w, step):
                    delta = rng.randint(-noise_amount, noise_amount)
                    for yy in range(y, min(y + step, h)):
                        for xx in range(x, min(x + step, w)):
                            v = pixels[xx, yy] + delta
                            pixels[xx, yy] = max(0, min(255, v))
            img = noisy

            # Slight rotation, expand canvas and fill with white so text
            # isn't cropped.
            angle = rng.uniform(-max_rotation_deg, max_rotation_deg)
            img = img.rotate(angle, expand=True, fillcolor=255, resample=Image.BICUBIC)

            # Mild blur to emulate scanner softness, then re-sharpen slightly
            # so OCR still has crisp edges -- net effect is "soft but legible".
            img = img.filter(ImageFilter.GaussianBlur(radius=0.4))

            # JPEG round-trip at the requested quality (this is what actually
            # bounds the per-page byte size).
            buf = io.BytesIO()
            img.convert("L").save(buf, format="JPEG", quality=jpeg_quality, optimize=True)
            buf.seek(0)
            jpeg_img = Image.open(buf)
            jpeg_img.load()

            # Re-embed as a full-page image PDF page via reportlab, at the
            # image's own aspect ratio scaled onto a US Letter page.
            page_buf = io.BytesIO()
            page_w, page_h = letter
            c = pdfcanvas.Canvas(page_buf, pagesize=letter)
            iw, ih = jpeg_img.size
            scale = min(page_w / iw, page_h / ih)
            draw_w, draw_h = iw * scale, ih * scale
            x_off = (page_w - draw_w) / 2
            y_off = (page_h - draw_h) / 2
            img_reader_buf = io.BytesIO()
            jpeg_img.save(img_reader_buf, format="JPEG", quality=jpeg_quality)
            img_reader_buf.seek(0)
            from reportlab.lib.utils import ImageReader
            c.drawImage(ImageReader(img_reader_buf), x_off, y_off, width=draw_w, height=draw_h)
            c.showPage()
            c.save()
            page_buf.seek(0)

            one_page_reader = PdfReader(page_buf)
            writer.add_page(one_page_reader.pages[0])

        with open(out_pdf_path, "wb") as f:
            writer.write(f)


# ---------------------------------------------------------------------------
# Photographed receipt / handwritten note
# ---------------------------------------------------------------------------

def _perspective_coeffs(src_quad, dst_quad):
    """Compute PIL PERSPECTIVE transform coefficients mapping dst_quad
    (the output corners) back to src_quad (the input corners)."""
    matrix = []
    for (x, y), (X, Y) in zip(dst_quad, src_quad):
        matrix.append([x, y, 1, 0, 0, 0, -X * x, -X * y])
        matrix.append([0, 0, 0, x, y, 1, -Y * x, -Y * y])
    A = matrix
    b = []
    for (X, Y) in src_quad:
        b.extend([X, Y])

    # Solve A @ coeffs = b via plain Gaussian elimination (avoid a numpy dep
    # here for a small 8x8 system).
    n = 8
    for col in range(n):
        pivot = max(range(col, n), key=lambda r: abs(A[r][col]))
        A[col], A[pivot] = A[pivot], A[col]
        b[col], b[pivot] = b[pivot], b[col]
        pv = A[col][col]
        for r in range(n):
            if r == col:
                continue
            factor = A[r][col] / pv
            for cIdx in range(col, n):
                A[r][cIdx] -= factor * A[col][cIdx]
            b[r] -= factor * b[col]
    return [b[i] / A[i][i] for i in range(n)]


def photograph_receipt(src_image_path: str, out_jpg_path: str, seed: int) -> None:
    """Take a flat rendered receipt/document image and make it look
    photographed: slight perspective warp, rotation, uneven lighting,
    JPEG artefacts. Deterministic given `seed`."""
    rng = random.Random(seed)
    os.makedirs(os.path.dirname(out_jpg_path) or ".", exist_ok=True)

    img = Image.open(src_image_path).convert("RGB")

    # Upscale small source renders before anything else touches them.
    #
    # The flat receipt renders come in around 660x810, which tesseract reads
    # as roughly 109 DPI -- far below the range where it resolves closed digit
    # counters reliably. The symptom is a corpus-wide scatter of single-digit
    # misreads on receipts that are perfectly legible by eye: $96.10 read back
    # as $98.10, $39,279.00 as $38,278.00, a date "21 Oct" as "21 02".
    # Softening the warp and blur did not touch it, because it is a pixel-size
    # problem, not a distortion problem -- upscaling the same images 3x before
    # OCR recovers every one of them exactly.
    #
    # A real phone photograph of a receipt is thousands of pixels wide, so
    # this is also the more faithful rendering. Done here rather than in each
    # generator so all three companies get it from one place.
    _MIN_WIDTH = 1900
    if img.width < _MIN_WIDTH:
        scale = _MIN_WIDTH / img.width
        img = img.resize(
            (int(img.width * scale), int(img.height * scale)), Image.LANCZOS
        )

    w, h = img.size

    # Pad canvas so perspective warp / rotation has margin to work with.
    pad = int(0.12 * max(w, h))
    canvas_img = Image.new("RGB", (w + 2 * pad, h + 2 * pad), (60, 58, 55))
    canvas_img.paste(img, (pad, pad))
    w2, h2 = canvas_img.size

    # Slight perspective: jitter each corner by up to 1.5% of dimension.
    # Kept deliberately modest: at 3% the combined warp/rotation/blur was
    # enough to cost tesseract whole lines (a counterparty name, a digit in a
    # total) on receipts that are plainly legible by eye. These documents
    # exist to be machine-read by the reader's own application, so an
    # unreadable figure makes the challenge impossible rather than hard.
    jitter = 0.015
    src_quad = [(0, 0), (w2, 0), (w2, h2), (0, h2)]
    dst_quad = [
        (rng.uniform(0, jitter * w2), rng.uniform(0, jitter * h2)),
        (w2 - rng.uniform(0, jitter * w2), rng.uniform(0, jitter * h2)),
        (w2 - rng.uniform(0, jitter * w2), h2 - rng.uniform(0, jitter * h2)),
        (rng.uniform(0, jitter * w2), h2 - rng.uniform(0, jitter * h2)),
    ]
    coeffs = _perspective_coeffs(src_quad, dst_quad)
    warped = canvas_img.transform((w2, h2), Image.PERSPECTIVE, coeffs, resample=Image.BICUBIC)

    # Slight rotation.
    angle = rng.uniform(-2.0, 2.0)
    rotated = warped.rotate(angle, expand=True, fillcolor=(60, 58, 55), resample=Image.BICUBIC)

    # Uneven lighting: a soft radial gradient multiplied over the image, with
    # the bright spot placed off-centre.
    rw, rh = rotated.size
    gradient = Image.new("L", (rw, rh), 0)
    gdraw = ImageDraw.Draw(gradient)
    cx, cy = rng.uniform(0.3, 0.7) * rw, rng.uniform(0.2, 0.6) * rh
    max_r = math.hypot(rw, rh)
    steps = 40
    for i in range(steps, 0, -1):
        frac = i / steps
        val = int(255 * (1 - 0.35 * (1 - frac)))
        r = max_r * frac
        gdraw.ellipse([cx - r, cy - r, cx + r, cy + r], fill=val)
    lit = Image.composite(rotated, Image.new("RGB", rotated.size, (0, 0, 0)), gradient)

    # Mild blur + noise for a phone-camera feel.
    lit = lit.filter(ImageFilter.GaussianBlur(radius=0.4))

    # Save at the highest quality that still fits the 500 KB per-page raster
    # cap. Stepping quality down beats shrinking the image: the upscale above
    # is what OCR needs, and JPEG quality degrades far more gracefully than
    # resolution does. In practice the first or second step is taken.
    for buf_quality in (88, 82, 76, 70, 64):
        lit.save(out_jpg_path, format="JPEG", quality=buf_quality, optimize=True)
        if os.path.getsize(out_jpg_path) <= 480_000:
            break


def handwritten_note_image(lines: list[str], out_jpg_path: str, seed: int,
                            width: int = 900, height: int = 600) -> None:
    """Render short text as a handwritten-*looking* cash receipt: per-word
    baseline jitter and rotation on a plain background, then photograph it.
    No handwriting font is assumed to be installed; jitter on a normal font
    is enough to read as informal/handwritten at receipt scale, and the
    photograph_receipt() pass (angle + lighting) does the rest."""
    rng = random.Random(seed)
    os.makedirs(os.path.dirname(out_jpg_path) or ".", exist_ok=True)

    img = Image.new("RGB", (width, height), (250, 248, 238))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans-Bold.ttf", 30)
    except OSError:
        font = ImageFont.load_default()

    y = 40
    for line in lines:
        x = 40 + rng.randint(-4, 4)
        for word in line.split(" "):
            jx = rng.randint(-2, 2)
            jy = rng.randint(-6, 6)
            jrot_layer = Image.new("RGBA", (300, 60), (0, 0, 0, 0))
            wdraw = ImageDraw.Draw(jrot_layer)
            wdraw.text((0, 0), word, font=font, fill=(20, 20, 40, 255))
            jrot_layer = jrot_layer.rotate(rng.uniform(-3, 3), expand=True, resample=Image.BICUBIC)
            img.paste(jrot_layer, (x + jx, y + jy), jrot_layer)
            bbox = draw.textbbox((0, 0), word + " ", font=font)
            x += (bbox[2] - bbox[0]) + 8
        y += 55

    tmp_flat = out_jpg_path + ".flat.png"
    img.save(tmp_flat)
    photograph_receipt(tmp_flat, out_jpg_path, seed=seed)
    os.remove(tmp_flat)


# ---------------------------------------------------------------------------
# Multi-document PDF concatenation
# ---------------------------------------------------------------------------

def concat_pdfs(pdf_paths: list[str], out_path: str) -> None:
    """Merge several unrelated PDFs into one multi-document PDF (e.g. Bright
    Harbor's 'march bills.pdf' containing several scanned vendor bills)."""
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    writer = PdfWriter()
    for p in pdf_paths:
        reader = PdfReader(p)
        for page in reader.pages:
            writer.add_page(page)
    with open(out_path, "wb") as f:
        writer.write(f)


# ---------------------------------------------------------------------------
# DOCX / XLSX / CSV
# ---------------------------------------------------------------------------

def render_docx(out_path: str, title: str, paragraphs: list[str],
                 table_rows: list[list[str]] | None = None) -> None:
    from docx import Document

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        table.style = "Table Grid"
        for row in table_rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    doc.save(out_path)


def render_xlsx(out_path: str, sheets: dict[str, list[list]]) -> None:
    """sheets: {sheet_name: [[row1 cells...], [row2 cells...], ...]}"""
    from openpyxl import Workbook

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name[:31])
        for row in rows:
            ws.append(row)
    wb.save(out_path)


def write_csv(out_path: str, rows: list[list], header: list[str] | None = None) -> None:
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        if header:
            w.writerow(header)
        w.writerows(rows)
